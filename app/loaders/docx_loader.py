"""
Loader cho file Word (.docx).
Trích xuất văn bản trực tiếp từ các paragraph và bảng, bỏ qua bước OCR.
Yêu cầu: python-docx (pip install python-docx)
"""

from docx import Document as DocxDocument
from pathlib import Path
from app.schemas.document import Document
from docx.oxml.ns import qn  # type: ignore
from docx.table import Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore


class DocxLoader:
    """Load file .docx thành DocumentSchema, không qua OCR."""

    def load(self, file_path: str | Path) -> Document:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        doc = DocxDocument(str(path))

        lines: list[str] = []

        for block in self._iter_block_items(doc):
            if hasattr(block, "text"):
                text = block.text.strip()
                if text:
                    lines.append(text)

        raw_text = "\n".join(lines).strip()

        # Lấy title từ core properties nếu có
        title = ""
        try:
            title = (doc.core_properties.title or "").strip()
        except Exception:
            pass
        if not title:
            title = path.stem

        return Document(
            doc_id=path.stem.lower().replace(" ", "_"),
            source_path=str(path),
            source_type="docx",
            title=title,
            raw_text=raw_text,
            metadata={
                "file_name": path.name,
                "paragraph_count": len(doc.paragraphs),
                "char_count": len(raw_text),
            },
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _iter_block_items(self, doc):
        """
        Yield từng paragraph và row của bảng theo thứ tự xuất hiện trong document body.
        Dùng docx.oxml để duyệt đúng thứ tự thay vì tách riêng paragraphs / tables.
        """

        body = doc.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                table = Table(child, doc)
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para
